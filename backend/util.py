from docx.shared import Pt
from docx.oxml.ns import qn

def add_work_experience(doc, work_exp):
    """
    Insert formatted work experience into a Word document.
    - Font: Garamond, size 10.5
    - First line: Company (bold), Location (normal) on right
    - Second line: Role (italic) on right, Dates on left
    - Achievements: bulleted list
    """
    if isinstance(work_exp, str):
        p = doc.add_paragraph(work_exp)
        for run in p.runs:
            run.font.name = "Garamond"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
            run.font.size = Pt(10.5)
        return

    for job in work_exp:
        # --- First line: Company (bold) on left, Location on right ---
        p1 = doc.add_paragraph()
        run_company = p1.add_run(job.get("company", ""))
        run_company.bold = True
        run_company.font.name = "Garamond"
        run_company._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
        run_company.font.size = Pt(10.5)

        # Tab for spacing
        p1.add_run("\t")

        run_location = p1.add_run(job.get("location", ""))
        run_location.font.name = "Garamond"
        run_location._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
        run_location.font.size = Pt(10.5)

        # --- Second line: Dates on left, Role (italic) on right ---
        p2 = doc.add_paragraph()
        run_dates = p2.add_run(job.get("dates", ""))
        run_dates.font.name = "Garamond"
        run_dates._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
        run_dates.font.size = Pt(10.5)

        p2.add_run("\t")

        run_role = p2.add_run(job.get("role", ""))
        run_role.italic = True
        run_role.font.name = "Garamond"
        run_role._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
        run_role.font.size = Pt(10.5)

        # --- Achievements as bullet points ---
        achievements = job.get("achievements", [])
        if isinstance(achievements, list):
            for a in achievements:
                if a.strip():
                    pa = doc.add_paragraph(style="List Bullet")
                    run_a = pa.add_run(a.strip())
                    run_a.font.name = "Garamond"
                    run_a._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
                    run_a.font.size = Pt(10.5)

        # Blank line between jobs
        doc.add_paragraph()




from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bullet_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Normal"  # fallback style
    
    # Apply bullet manually
    p_pr = p._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), "0")
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    p_pr.append(numPr)

    for run in p.runs:
        run.font.name = "Garamond"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
        run.font.size = Pt(10.5)

    return p
