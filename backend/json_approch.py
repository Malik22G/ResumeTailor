import json
import re
import mammoth
from bs4 import BeautifulSoup
from docx import Document
from groq import Groq
from config import GROQ_API_KEY
from docx.shared import Pt
from docx.oxml.ns import qn

from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize Groq client
client = Groq(api_key=GROQ_API_KEY)

from docx import Document

doc = Document("data/resume_template.docx")

for style in doc.styles:
    print(style.name)

# ==========================
# STEP 1: Parse User Resume
# ==========================
def parse_resume_sections(input_docx):
    """Extract basic sections from the user's resume into a dictionary."""
    with open(input_docx, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value

    soup = BeautifulSoup(html, "html.parser")
    text = soup.get_text("\n").split("\n")

    sections = {
        "summary": "",
        "education": "",
        "work_experience": "",
        "projects": "",
        "skills": "",
        "additional": "",
        "activities": ""
    }

    current_section = None
    for line in text:
        line = line.strip()
        if not line:
            continue

        lower = line.lower()
        if "summary" in lower:
            current_section = "summary"
            continue
        elif "education" in lower:
            current_section = "education"
            continue
        elif "work experience" in lower:
            current_section = "work_experience"
            continue
        elif "project" in lower:
            current_section = "projects"
            continue
        elif "skill" in lower:
            current_section = "skills"
            continue
        elif "additional" in lower:
            current_section = "additional"
            continue
        elif "activit" in lower:
            current_section = "activities"
            continue

        if current_section:
            sections[current_section] += line + "\n"

    return sections

# ==========================
# STEP 2: LLM Rewrite
# ==========================
def rewrite_resume_sections(sections, job_description):
    """Ask LLM to rewrite ALL resume sections."""
    prompt = f"""
You are a professional resume writer.
Your task is to update ALL of the candidate's resume sections to better match the job description.

Input:
- Candidate resume data is provided in JSON format.
- The JSON contains these keys: {list(sections.keys())}

Instructions:
1. Rewrite every section: summary, education, work_experience, projects, skills, additional, activities.
2. Make each section concise, professional, and tailored to the job description.
3. Emphasize relevant keywords from the job description.
4. Quantify achievements where possible.
5. Do not invent jobs, degrees, or skills not supported by the candidate’s experience.
6. Return ONLY valid JSON, with the SAME keys as the input.

Candidate Resume Data:
{json.dumps(sections, indent=2)}

Job Description:
{job_description}
"""

    response = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.5,
        max_tokens=1500,
    )

    text = response.choices[0].message.content.strip()

    # Extract JSON safely
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("LLM did not return valid JSON")

    data = json.loads(match.group(0))
    return data

# ==========================
# STEP 3: Fill Template
# ==========================
def format_work_experience(work_exp):
    """Format work experience with bullets if JSON structured, otherwise keep string."""
    if isinstance(work_exp, str):
        return work_exp
    
    formatted = []
    for job in work_exp:
        header = f"{job.get('company','')}  {job.get('location','')}"
        role = f"{job.get("role", "")} \t\t{job.get('dates','')}"

        formatted.append(header)
        if role:
            formatted.append(role)
        
        achievements = job.get("achievements", "")

        if isinstance(achievements, list):
            for a in achievements:
                if a.strip():
                    formatted.append(f"• {a.strip()}")
        elif isinstance(achievements, str):  # fallback if it's a string
            for a in achievements.split("\n"):
                if a.strip():
                    formatted.append(f"• {a.strip()}")

        formatted.append("")  # blank line between jobs
    
    return "\n".join(formatted)



def set_run_font(run, font_name="Garamond", size=Pt(10.5), bold=False, italic=False):
    """Forcefully set font for a run (Word-compatible)."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)  # important!
    run.font.size = size
    run.bold = bold
    run.italic = italic


def insert_formatted_experience(doc, placeholder, work_exp):
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = ""  # clear placeholder text

            for job in work_exp:
                # Company + Location
                header = p.insert_paragraph_before(f"{job.get('company','')}  {job.get('location','')}")
                set_run_font(header.runs[0], bold=True)

                # Role + Dates
                role_text = f"{job.get('role','')}   {job.get('dates','')}"
                if role_text.strip():
                    role_para = p.insert_paragraph_before(role_text)
                    set_run_font(role_para.runs[0], italic=True)

                # Achievements with manual bullets
                achievements = job.get("achievements", [])
                if isinstance(achievements, str):
                    achievements = [a.strip() for a in achievements.split("\n") if a.strip()]

                for a in achievements:
                    bullet_para = p.insert_paragraph_before(f"• {a}")
                    set_run_font(bullet_para.runs[0])

            break

    """Insert work experience into the document with Garamond font and 10.5 pt size."""
    for p in doc.paragraphs:
        if placeholder in p.text:
            p.text = ""  # clear placeholder text

            for job in work_exp:
                # Company + Location
                header = p.insert_paragraph_before(f"{job.get('company','')}  {job.get('location','')}")
                set_run_font(header.runs[0], bold=True)

                # Role + Dates
                role_text = f"{job.get('role','')}   {job.get('dates','')}"
                if role_text.strip():
                    role_para = p.insert_paragraph_before(role_text)
                    set_run_font(role_para.runs[0], italic=True)

                # Achievements as bullets
                achievements = job.get("achievements", [])
                if isinstance(achievements, str):
                    achievements = [a.strip() for a in achievements.split("\n") if a.strip()]

                for a in achievements:
                    bullet = p.insert_paragraph_before(a, style="Bullet")
                    set_run_font(bullet.runs[0])

            break  # stop once placeholder is replaced

def fill_resume_template(template_path, output_path, rewritten_sections):
    """Replace placeholders in the template with rewritten resume content."""
    doc = Document(template_path)

    def normalize(val, section=""):
        if section in ("work_experience","projects") :
            return format_work_experience(val)
        if isinstance(val, list):
            return "\n".join(f"• {v}" for v in val)
        return str(val)
    print(rewritten_sections["work_experience"])
    placeholder_map = {
        "{{SUMMARY}}": normalize(rewritten_sections.get("summary", "")),
        "{{EDUCATION}}": normalize(rewritten_sections.get("education", "")),
        #"{{WORK_EXPERIENCE}}": normalize(rewritten_sections.get("work_experience", ""), "work_experience"),
        "{{PROJECTS}}": normalize(rewritten_sections.get("projects", ""),"projects"),
        "{{SKILLS}}": normalize(rewritten_sections.get("skills", "")),
        "{{ADDITIONAL}}": normalize(rewritten_sections.get("additional", "")),
        "{{ACTIVITIES}}": normalize(rewritten_sections.get("activities", "")),
    }

    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, val in placeholder_map.items():
            if key in p.text:
                for run in p.runs:
                    run.text = ""
                p.add_run(val)

    # Special handling for work experience (styled insertion)
    if isinstance(rewritten_sections.get("work_experience"), list):
        insert_formatted_experience(doc, "{{WORK_EXPERIENCE}}", rewritten_sections["work_experience"])
    else:
        # fallback to plain string if it's not structured
        for p in doc.paragraphs:
            if "{{WORK_EXPERIENCE}}" in p.text:
                p.text = rewritten_sections["work_experience"]

    doc.save(output_path)

# ==========================
# MAIN
# ==========================
if __name__ == "__main__":
    input_resume = "data/resume.docx"
    job_desc_file = "data/job_desc.txt"
    template_resume = "data/resume_template.docx"
    output_resume = "results/tailored_resume.docx"

    print("Parsing user resume...")
    sections = parse_resume_sections(input_resume)

    print("Reading job description...")
    with open(job_desc_file, "r", encoding="utf-8") as f:
        job_description = f.read()

    print("Requesting LLM rewrite...")
    rewritten = rewrite_resume_sections(sections, job_description)

    print("Filling template...")
    fill_resume_template(template_resume, output_resume, rewritten)

    print(f"✅ Tailored resume saved to: {output_resume}")
